import shutil, os
import random
from googletrans import Translator
from gtts import gTTS
import PyPDF2
from flask import Flask, request, render_template,send_file
import docx2txt
import docx
import pytesseract



app = Flask(__name__)

# app = Flask(__name__, static_folder="images")
nrt=0
ntr=0
l="kn"
APP_ROOT = os.path.dirname(os.path.abspath(__file__))
word="pdf"
word2='docx'
hit=""
n=""
filename2 = ""
def isWordPresent(sentence, word):
    # To break the sentence in words
    s = sentence.split(".")

    for i in s:

        # Comparing the current word
        # with the word to be searched
        if (i == word):
            return True
    return False
def isWordPresent2(sentence, word2):
    # To break the sentence in words
    s = sentence.split(".")

    for i in s:

        # Comparing the current word
        # with the word to be searched
        if (i == word2):
            return True
    return False



@app.route("/")
def index():
    return render_template("screen1.html")
@app.route("/screen1")
def index11():
    return render_template("primary.html")
@app.route("/screen2")
def index1():
    return render_template("screen2.html")
@app.route("/2screen2")
def index2():
    return render_template("2screen2.html")
@app.route("/4screen2")
def index3():
    return render_template("4screen2.html")

@app.route("/translate", methods=["POST"])
def tran():
    htts=request.form['htts']
    des = request.form['id']
    des2 = request.form['id2']
    print(htts)
    trans = Translator()
    if htts =="":
        return render_template('fake3.html')
    else:
        try:
            t = trans.translate(htts, src=des, dest=des2)
            return render_template('primary.html',variable=t.text)
        except:
            return render_template("fake3.html")
@app.route("/upload", methods=["POST"])
def upload():
    global nrt
    nrt = random.randint(1, 10000000)
    nrt2 = str(nrt)
    nrt3 = nrt2 + "fld/"
    i = 1
    global n
    n=request.form['id']
    global hit
    hit=request.form['id2']
    if os.path.exists('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt3+'/comp.pdf'):
        os.remove('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt3+'/comp.pdf')
    else:
        os.mkdir('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt3)
    target = os.path.join(APP_ROOT,"fld/"+nrt2)
    # target = os.path.join(APP_ROOT, 'static/')
    print(target)
    if not os.path.isdir(target):
            os.mkdir(target)
    else:
        print("Couldn't create upload directory: {}".format(target))
    print(request.files.getlist("file"))
    for upload in request.files.getlist("file"):
        print(upload)
        print("{} is the file name".format(upload.filename))
        if upload.filename == "":
            return render_template("fake.html")
        else:
            pass
        if (isWordPresent(upload.filename, word)):
            print("yes")
            pass
        else:
           return render_template("fake.html")
        global filename2
        filename2=upload.filename
        filename ="comp.pdf"
        destination = "/".join([target, filename])
        print ("Accept incoming file:", filename)
        print ("Save it to:", destination)
        upload.save(destination)
        os.chdir('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt2)
        try:
            with open('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt2+'/comp.pdf', 'rb') as pdf_file,open('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt2+'/comp.txt', 'w', encoding="utf-8") as text_file:
                read_pdf = PyPDF2.PdfFileReader(pdf_file)
                number_of_pages = read_pdf.getNumPages()
                for page_number in range(number_of_pages):  # use xrange in Py2
                    page = read_pdf.getPage(page_number)
                    page_content = page.extractText()
                    trans = Translator()
                    try:
                        t = trans.translate(page_content, src='en', dest=n)
                        print(str(t.text))
                        text_file.write(t.text)
                    except:
                        pass
            text_file.close()
        except:
            return render_template("fake.html")

    if hit==".mp3":
        k="screen3.html"
    else:
        k="screen4.html"
    os.remove('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/' + nrt2 + '/comp.pdf')
    files = ['comp.txt']
    for f in files:
        shutil.move(f, 'C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt3)
    path = "C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/" + nrt3 + "/comp.txt"
    fios = open(path, 'r', encoding="utf-8")
    fios2 = fios.read()
    print(hit)
    try:
        shutil.rmtree('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/')
    except:
        pass
    if hit == ".docx":
        print("came here")
        doc = docx.Document()
        doc.add_paragraph(fios2)
        doc.save('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/' + nrt3 +'/'+filename2+'.docx')
        h = 'C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/' + nrt3 + '/helloworld.docx'
    elif hit == ".mp3":
        print("came here")
        toSpeak = gTTS(text=fios2, lang=n, slow=False)
        file = 'C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/' + nrt3 +'/'+filename2+ '.mp3'
        toSpeak.save(file)
        print("sucess")
    return render_template(k, name=filename)

@app.route('/upload')
def send_image():
    print("loading")
    nrt2=str(nrt)
    nrt3 = nrt2 + "fld/"
    load='C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt3+'/'+filename2+hit
    return send_file(load,as_attachment=True)

@app.route("/mean2", methods=["POST"])
def upload2():
    import docx
    text = []
    global nrt
    nrt = random.randint(1, 10000000)
    nrt2 = str(nrt)
    nrt3 = nrt2 + "fld/"
    i = 1
    global n
    n=request.form['id']
    global hit
    hit=request.form['id2']
    if os.path.exists('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt3+'/comp.pdf'):
        os.remove('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt3+'/comp.pdf')
    else:
        os.mkdir('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt3)
    target = os.path.join(APP_ROOT,"fld/"+nrt2)
    # target = os.path.join(APP_ROOT, 'static/')
    print(target)
    if not os.path.isdir(target):
            os.mkdir(target)
    else:
        print("Couldn't create upload directory: {}".format(target))
    print(request.files.getlist("file"))
    for upload in request.files.getlist("file"):
        print(upload)
        print("{} is the file name".format(upload.filename))
        if upload.filename == "":
            return render_template("fake.html")
        else:
            pass
        if (isWordPresent2(upload.filename,word2)):
            print("yes")
            pass
        else:
           return render_template("fake2.html")
        global filename2
        filename2 = upload.filename

        filename ="comp.docx"
        destination = "/".join([target, filename])
        print ("Accept incoming file:", filename)
        print ("Save it to:", destination)
        upload.save(destination)
        os.chdir('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt2)
        doc = docx.Document('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt2+'/comp.docx')
        text_file=open('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt2+'/comp.txt', 'w', encoding="utf-8")
        try:
            for para in doc.paragraphs:
                text.append(para.text)
                trans = Translator()
                try:
                    t = trans.translate(para.text, src='en', dest=n)
                    print(f'{t.text}')
                    text_file.write(t.text)
                except:
                    pass
        except:
            return render_template("fake.html")

        text_file.close()

    if hit==".mp3":
        k="2screen4.html"
    else:
        k="2screen3.html"
    os.remove('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/' + nrt2 + '/comp.docx')
    files = ['comp.txt']
    for f in files:
        shutil.move(f, 'C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt3)
    path = "C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/" + nrt3 + "/comp.txt"
    fios = open(path, 'r', encoding="utf-8")
    fios2 = fios.read()
    print(hit)
    try:
        shutil.rmtree('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/')
    except:
        pass
    if hit == ".docx":
        doc = docx.Document()
        doc.add_paragraph(fios2)
        doc.save('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/' + nrt3 +'/'+filename2+'.docx')
        h = 'C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/' + nrt3 + '/helloworld.docx'
    elif hit == ".mp3":
        toSpeak = gTTS(text=fios2, lang=n, slow=False)
        file = 'C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/' + nrt3 +'/'+filename2+'.mp3'
        toSpeak.save(file)
        print("sucess")
    return render_template(k, name=filename)


@app.route('/mean2')
def send_image2():
    nrt2=str(nrt)
    nrt3 = nrt2 + "fld/"
    load='C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+ nrt3 +'/'+filename2+hit
    return send_file(load,as_attachment=True)


@app.route("/mean3", methods=["POST"])
def upload3():
    global nrt
    nrt = random.randint(1, 10000000)
    nrt2 = str(nrt)
    nrt3 = nrt2 + "fld/"
    i = 1
    global n
    n=request.form['id']
    global hit
    hit=request.form['id2']
    if os.path.exists('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt3+'/comp.pdf'):
        os.remove('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt3+'/comp.pdf')
    else:
        os.mkdir('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt3)
    target = os.path.join(APP_ROOT,"fld/"+nrt2)
    # target = os.path.join(APP_ROOT, 'static/')
    print(target)
    if not os.path.isdir(target):
            os.mkdir(target)
    else:
        print("Couldn't create upload directory: {}".format(target))
    print(request.files.getlist("file"))
    for upload in request.files.getlist("file"):
        print(upload)
        print("{} is the file name".format(upload.filename))
        if upload.filename == "":
            print("here")
            return render_template("fake.html")
        else:
            pass
        filename =upload.filename
        destination = "/".join([target, filename])
        print ("Accept incoming file:", filename)
        print ("Save it to:", destination)
        upload.save(destination)
        os.chdir('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt2)
        try:
            text_file = open('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/' + nrt2 + '/comp.txt', 'w',encoding="utf-8")
            im = ('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt2+'/'+filename)
            config = ('-l eng --oem 1 --psm 3')
            text = pytesseract.image_to_string(im, config=config)
            print(text)
            trans = Translator()
            t = trans.translate(text, src='en', dest=n)
            print(f'{t.text}')
            text_file.write(t.text)
            text_file.close()
        except:
            return render_template("fake.html")
        global filename2
        filename2 = upload.filename

    if hit==".mp3":
        k="4screen3.html"
    else:
        k="4screen4.html"
    files = ['comp.txt']
    for f in files:
        shutil.move(f, 'C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+nrt3)
    path = "C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/" + nrt3 + "/comp.txt"
    fios = open(path, 'r', encoding="utf-8")
    fios2 = fios.read()
    print(hit)
    try:
        shutil.rmtree('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/')
        shutil.rmtree('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/')
    except:
        pass
    if hit == ".docx":
        doc = docx.Document()
        doc.add_paragraph(fios2)
        doc.save('C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/' + nrt3 +'/'+filename2+'.docx')
        h = 'C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/' + nrt3 + '/helloworld.docx'
    elif hit == ".mp3":
        toSpeak = gTTS(text=fios2, lang=n, slow=False)
        file = 'C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+ nrt3 +'/'+filename2+'.mp3'
        toSpeak.save(file)
        print("sucess")
    return render_template(k, name=filename)


@app.route('/mean3')
def send_image3():
    nrt2=str(nrt)
    nrt3 = nrt2 + "fld/"
    load='C:/Users/vishnu sai/Desktop/my projectes/pdf/fld/'+ nrt3 +'/'+filename2+hit
    return send_file(load,as_attachment=True)

if __name__ == "__main__":
    app.run(port=4555, debug=True)